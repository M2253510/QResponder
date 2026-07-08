"""Format-perfect write-back into the ORIGINAL file (§15) — Phase 2.

Covers every ingested questionnaire format: xlsx/xlsm and docx fill answers into
their native cells/anchors; pdf preserves the original pages and appends a
grounded 'Responses' section (arbitrary PDFs have no editable cell model).

Fills each answer into the user's own template, in a COPY, honoring the traps
that silently corrupt files:
  * merged ranges — write only to the top-left anchor (writing to a non-anchor
    MergedCell raises); never to a merged member cell;
  * shared/immutable styles — set `.value` only, never mutate a shared style;
  * embedded images/charts — openpyxl may drop them on save, so if the workbook
    has any, we DON'T save over them: we warn and fall back to the separate
    Phase-0/1 output file rather than stripping the user's diagrams;
  * the original is never overwritten — output goes to `<name>.answered.xlsx`.

Only ANSWERED results (incl. resolved attachments, referenced by filename) are
written; NEEDS_REVIEW cells are left blank for the human.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

from ..models import QuestionnaireResult, Status

log = logging.getLogger("qresponder.writeback")

_ANSWER_HEADER_RE = re.compile(r"\b(answer|response|comment|comments|reply)\b", re.IGNORECASE)


def _cell_value(r, review_markers: bool = True) -> str | None:
    """The value to write, or None to skip. ANSWERED → the answer; NEEDS_REVIEW →
    a visible marker (Phase 7 C) so unresolved cells aren't silently blank."""
    if r.status == Status.ANSWERED:
        if r.attachment_path:
            return r.answer or Path(r.attachment_path).name
        return r.answer or None
    if review_markers:
        from .writer import review_marker

        return review_marker(r)
    return None  # markers off → leave NEEDS_REVIEW blank (legacy behavior)


# --- xlsx --------------------------------------------------------------------

def _has_media(wb) -> bool:
    for ws in wb.worksheets:
        if getattr(ws, "_images", None) or getattr(ws, "_charts", None):
            return True
    return False


def _merged_anchor(ws, coord: str) -> str:
    """If coord lies inside a merged range, return the range's top-left anchor."""
    from openpyxl.utils import get_column_letter

    for mrange in ws.merged_cells.ranges:
        if coord in mrange:
            return f"{get_column_letter(mrange.min_col)}{mrange.min_row}"
    return coord


def _parse_sheet_coord(anchor: str, default_ws):
    """'Sheet!C7' -> (worksheet, 'C7'); 'C7' -> (default_ws, 'C7')."""
    if "!" in anchor:
        sheet, coord = anchor.split("!", 1)
        return sheet, coord
    return (default_ws.title if default_ws else None), anchor


def _resolve_xlsx_target(wb, r):
    """Return (worksheet, coord) for where the answer goes, or None."""
    from openpyxl.utils import column_index_from_string, get_column_letter
    from openpyxl.utils.cell import coordinate_from_string

    # 1. Explicit answer anchor.
    if r.answer_location_hint:
        sheet, coord = _parse_sheet_coord(r.answer_location_hint, wb.active)
        ws = wb[sheet] if sheet in wb.sheetnames else wb.active
        return ws, coord

    # 2. Heuristic from the question's own cell.
    if r.location_hint:
        sheet, coord = _parse_sheet_coord(r.location_hint, wb.active)
        ws = wb[sheet] if sheet in wb.sheetnames else wb.active
        try:
            col_letter, row = coordinate_from_string(coord)
            col = column_index_from_string(col_letter)
        except Exception:  # noqa: BLE001
            return None
        # 2a. A column whose header matches answer|response|comment — but only
        # if that row's cell is empty (never overwrite a pre-filled value, SH3).
        header_coord = None
        for header_row in (1, 2):
            for c in ws[header_row]:
                if c.value and _ANSWER_HEADER_RE.search(str(c.value)):
                    header_coord = f"{get_column_letter(c.column)}{row}"
                    break
            if header_coord:
                break
        if header_coord and ws[_merged_anchor(ws, header_coord)].value in (None, ""):
            return ws, header_coord
        # 2b. First empty cell to the right in the question's row.
        for delta in (1, 2, 3):
            cand = f"{get_column_letter(col + delta)}{row}"
            if ws[_merged_anchor(ws, cand)].value in (None, ""):
                return ws, cand
    return None


def _list_options(formula1: str) -> list[str]:
    """Parse a list data-validation's options from formula1 (e.g. '"Yes,No"')."""
    if not formula1:
        return []
    f = formula1.strip()
    if f.startswith('"') and f.endswith('"'):
        f = f[1:-1]
    return [o.strip() for o in f.split(",") if o.strip()]


def _coerce_to_validation(ws, coord: str, value: str) -> str:
    """If `coord` has a list dropdown validation and `value` isn't an allowed
    option, map it to one (e.g. 'Yes. ...' -> 'Yes'); else return value unchanged.
    Never modifies the validation object."""
    try:
        for dv in ws.data_validations.dataValidation:
            if dv.type != "list" or coord not in dv.sqref:
                continue
            options = _list_options(dv.formula1 or "")
            if not options or value in options:
                return value
            vlow = value.lower()
            # Prefer a word-boundary-ish match (yes/no), else substring.
            for opt in options:
                if vlow == opt.lower() or vlow.startswith(opt.lower()):
                    return opt
            for opt in options:
                if opt.lower() in vlow:
                    return opt
            return value
    except Exception:  # noqa: BLE001 - never let validation handling break write-back
        return value
    return value


def _writeback_xlsx(result: QuestionnaireResult, original_path: Path, out_path: Path,
                    review_markers: bool = True) -> dict:
    import openpyxl
    from openpyxl.cell.cell import MergedCell

    wb = openpyxl.load_workbook(original_path)
    if _has_media(wb):
        wb.close()
        log.warning(
            "%s contains images/charts; openpyxl may drop them on save. "
            "Skipping write-back; use the separate answered output instead.",
            original_path.name,
        )
        return {"written": None, "fallback": True, "reason": "workbook has images/charts"}

    written = 0
    for r in result.results:
        value = _cell_value(r, review_markers=review_markers)
        if value is None:
            continue
        target = _resolve_xlsx_target(wb, r)
        if target is None:
            continue
        ws, coord = target
        coord = _merged_anchor(ws, coord)  # never write to a non-anchor merged cell
        cell = ws[coord]
        if isinstance(cell, MergedCell):  # defensive: anchor resolution failed
            continue
        # Universal guard: never overwrite a pre-filled cell (applies to explicit
        # answer anchors too — markers and answers alike).
        if cell.value not in (None, ""):
            continue
        # If the cell has a list/dropdown data-validation, write an ALLOWED value
        # when we can map to one — and never touch the validation object itself
        # (openpyxl preserves data validations across load/save; we only set the
        # value, so dropdowns/validations survive — Part F).
        value = _coerce_to_validation(ws, coord, value)
        cell.value = value  # set value only; never touch the (shared) style/validation
        # Red-italic the marker in unresolved cells (Phase 11 F). Only ever applied
        # to a cell we just filled with a marker (it was blank) — ANSWERED cells and
        # validations/media are untouched. Best-effort: never break write-back.
        if review_markers and r.status == Status.NEEDS_REVIEW:
            try:
                from openpyxl.styles import Font

                cell.font = Font(color="FFC0392B", italic=True)
            except Exception:  # noqa: BLE001
                pass
        written += 1

    wb.save(out_path)
    wb.close()
    return {"written": str(out_path), "fallback": False, "cells": written}


# --- docx --------------------------------------------------------------------

def _writeback_docx(result: QuestionnaireResult, original_path: Path, out_path: Path,
                    review_markers: bool = True) -> dict:
    import docx as _docx

    document = _docx.Document(str(original_path))
    paras = document.paragraphs
    tables = document.tables
    written = 0
    for r in result.results:
        value = _cell_value(r, review_markers=review_markers)
        anchor = r.answer_location_hint
        if value is None or not anchor:
            continue
        m_para = re.match(r"^para\[(\d+)\]$", anchor)
        m_cell = re.match(r"^table\[(\d+)\]\.r(\d+)\.c(\d+)$", anchor)
        if m_para:
            idx = int(m_para.group(1))
            if 0 <= idx < len(paras):
                paras[idx].add_run(" " + value)
                written += 1
        elif m_cell:
            ti, ri, ci = (int(x) for x in m_cell.groups())
            if ti < len(tables) and ri < len(tables[ti].rows) and ci < len(tables[ti].rows[ri].cells):
                tables[ti].rows[ri].cells[ci].text = value
                written += 1
    document.save(out_path)
    return {"written": str(out_path), "fallback": False, "cells": written}


# --- pdf ----------------------------------------------------------------------

def _pdf_txt(s: str) -> str:
    """fpdf2's core fonts are Latin-1; map anything outside it (smart quotes,
    em-dashes, bullets) to a safe representation so answer text never crashes the
    writer."""
    subs = {"‘": "'", "’": "'", "“": '"', "”": '"',
            "–": "-", "—": "-", "•": "-", "…": "...", " ": " "}
    for k, v in subs.items():
        s = s.replace(k, v)
    return s.encode("latin-1", "replace").decode("latin-1")


def _writeback_pdf(result: QuestionnaireResult, original_path: Path, out_path: Path,
                   review_markers: bool = True) -> dict:
    """PDF write-back. Arbitrary PDFs aren't reliably editable in place (fixed
    layouts, no cell model), so we PRESERVE the original pages untouched and
    APPEND a clearly-labeled 'Responses' section — each question with its grounded
    answer (or a review marker) and source. The original is never overwritten."""
    try:
        from io import BytesIO

        from fpdf import FPDF
        from fpdf.enums import XPos, YPos
        from pypdf import PdfReader, PdfWriter
    except ImportError as exc:  # pragma: no cover - deps are in base install
        return {"written": None, "fallback": True,
                "reason": f"PDF write-back needs pypdf + fpdf2 ({exc})"}

    # Each multi_cell returns the cursor to the left margin and drops a line, so a
    # width-0 (full-width) cell always has room to render.
    def _cell(size, text, style="", rgb=(0, 0, 0), h=5.5):
        pdf.set_font("Helvetica", style, size)
        pdf.set_text_color(*rgb)
        pdf.multi_cell(0, h, _pdf_txt(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(15, 15, 15)
    pdf.add_page()
    _cell(15, "QRESPONDER - Responses", style="B", h=8)
    _cell(8, f"Answers for {original_path.name}. Grounded in your knowledge base; "
             "review any flagged items before sending. The original questionnaire "
             "pages precede this section, unchanged.", rgb=(110, 110, 110), h=5)
    pdf.ln(3)

    written = 0
    for i, r in enumerate(result.results, start=1):
        value = _cell_value(r, review_markers=review_markers)
        if value is None:
            continue
        _cell(10, f"{i}. {r.question_text}", style="B")
        needs = r.status == Status.NEEDS_REVIEW
        _cell(10, value, style="I" if needs else "",
              rgb=(192, 57, 43) if needs else (0, 0, 0))  # match xlsx/docx marker red
        srcs = ", ".join(dict.fromkeys(c.source for c in r.citations if c.source))
        if srcs:
            _cell(8, f"Source: {srcs}", rgb=(110, 110, 110), h=4.5)
        pdf.ln(3)
        written += 1

    answers_pdf = BytesIO(bytes(pdf.output()))

    writer = PdfWriter()
    original_merged = True
    try:
        writer.append(PdfReader(str(original_path)))  # original pages, untouched
    except Exception as exc:  # noqa: BLE001 - corrupt/locked source: emit answers alone
        log.warning("Could not merge original PDF (%s); writing answers-only PDF.", exc)
        writer = PdfWriter()
        original_merged = False
    writer.append(PdfReader(answers_pdf))
    with open(out_path, "wb") as f:
        writer.write(f)
    writer.close()
    return {"written": str(out_path), "fallback": False,
            "answers": written, "original_preserved": original_merged}


# --- entry point -------------------------------------------------------------

def has_answer_anchors(result: QuestionnaireResult) -> bool:
    return any(r.answer_location_hint for r in result.results)


def write_back(result: QuestionnaireResult, original_path: str, out_dir: str,
               review_markers: bool = True) -> dict:
    """Write answers into a COPY of the original template. Returns a status dict
    with 'written' (path or None) and 'fallback' (bool)."""
    src = Path(original_path)
    out = Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)
    ext = src.suffix.lower()
    out_path = out / f"{src.stem}.answered{ext}"

    if ext in {".xlsx", ".xlsm"}:
        return _writeback_xlsx(result, src, out_path, review_markers=review_markers)
    if ext == ".docx":
        return _writeback_docx(result, src, out_path, review_markers=review_markers)
    if ext == ".pdf":
        return _writeback_pdf(result, src, out_path, review_markers=review_markers)
    return {"written": None, "fallback": True, "reason": f"write-back unsupported for {ext}"}
